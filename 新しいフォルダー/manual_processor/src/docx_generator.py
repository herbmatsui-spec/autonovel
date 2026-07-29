"""
Word Document Generator Module
Creates formatted Word documents from processed text
"""

import logging
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


def clean_markdown_bold(text: str) -> str:
    """
    Remove markdown bold markers (**) from text.
    
    Args:
        text: Text that may contain **bold** markdown
        
    Returns:
        Text with ** markers removed
    """
    if not text:
        return ""
    # Remove **bold** markers - replace **text** with text
    # Handle both **text** and ** text ** patterns
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    return text


class WordGenerator:
    """Generates formatted Word documents"""
    
    def __init__(self):
        """Initialize Word generator"""
        if Document is None:
            raise ImportError("python-docx is required for Word generation. Install with: pip install python-docx")
        logger.debug("WordGenerator initialized")
    
    def generate_word(self, content: str, output_path: Path, title: str = "Processed Manual", compact_layout: bool = False, use_emojis: bool = False, remove_markdown_bold: bool = True, qr_code_path: Optional[Path] = None, qr_audio_url: Optional[str] = None) -> Path:
        """
        Generate a formatted Word document with optional QR code for audio playback
        
        Args:
            content: Text content for the Word document
            output_path: Path for the output file
            title: Document title
            compact_layout: If True, reduce spacing and increase text size for compact output
            use_emojis: If True, insert emojis into headings and bullets
            remove_markdown_bold: If True, remove ** markdown bold markers from text
            qr_code_path: Optional path to QR code image
            qr_audio_url: Optional URL for audio playback
            
        Returns:
            Path to generated Word file
        """
        from docx.shared import Inches
        
        # Remove markdown bold markers if requested
        if remove_markdown_bold:
            content = clean_markdown_bold(content)
        
        document = Document()
        
        # Title
        display_title = f"📄 {title}" if use_emojis else title
        heading = document.add_heading(display_title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if compact_layout:
            for run in heading.runs:
                run.font.size = Pt(22)
        
        # Add QR Code Callout Box if QR code image is provided
        if qr_code_path and qr_code_path.exists():
            try:
                table = document.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.autofit = False
                
                # Left cell: QR Code Image
                cell_qr = table.cell(0, 0)
                cell_qr.width = Inches(1.8)
                p_qr = cell_qr.paragraphs[0]
                p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_qr = p_qr.add_run()
                run_qr.add_picture(str(qr_code_path), width=Inches(1.5))
                
                # Right cell: Explanatory text
                cell_text = table.cell(0, 1)
                cell_text.width = Inches(4.5)
                p_text = cell_text.paragraphs[0]
                run_title = p_text.add_run("📱 スマートフォンで聴ける音声解説\n")
                run_title.font.bold = True
                run_title.font.size = Pt(12)
                
                run_desc = p_text.add_run("スマホのカメラで上のQRコードを読み取ると、このマニュアルの要点音声解説を再生できます。研修・作業中の確認にご活用ください。")
                run_desc.font.size = Pt(10)
                
                if qr_audio_url and qr_audio_url.startswith("http"):
                    p_url = cell_text.add_paragraph()
                    run_url = p_url.add_run(f"Web URL: {qr_audio_url}")
                    run_url.font.size = Pt(8.5)
                
                document.add_paragraph()  # Spacing
            except Exception as e:
                logger.warning(f"Could not embed QR code in Word document: {e}")
        
        base_font_size = Pt((13 if compact_layout else 11))
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Check if this line is part of a markdown table
            if line.strip().startswith('|') and line.strip().endswith('|') and len(line.strip().split('|')) > 2:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|') and len(lines[i].strip().split('|')) > 2:
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Render table
                rows_data = []
                for tline in table_lines:
                    # Skip separator line like |---|---|
                    if re.match(r'^\|[\s:-|-]+\|$', tline):
                        continue
                    cells = [c.strip() for c in tline.split('|')[1:-1]]
                    if any(cells):
                        rows_data.append(cells)
                
                if rows_data:
                    num_rows = len(rows_data)
                    num_cols = max(len(row) for row in rows_data)
                    table = document.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for r_idx, row in enumerate(rows_data):
                        for c_idx, cell_value in enumerate(row):
                            if c_idx < num_cols:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = cell_value
                                # Style header row
                                if r_idx == 0:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.font.bold = True
                    # Add paragraph spacing after table
                    document.add_paragraph()
                continue

            if line.strip().startswith('## '):
                # Section header
                header_text = line[3:].strip()
                if use_emojis:
                    header_text = f"📌 {header_text}"
                heading = document.add_heading(header_text, level=(1 if not compact_layout else 1))
                if compact_layout:
                    for run in heading.runs:
                        run.font.size = Pt(16)
                        heading.paragraph_format.space_after = Pt(2)
                        heading.paragraph_format.space_before = Pt(2)
            elif line.strip().startswith('### '):
                # Sub-section header
                header_text = line[4:].strip()
                if use_emojis:
                    header_text = f"📌 {header_text}"
                heading = document.add_heading(header_text, level=2)
                if compact_layout:
                    for run in heading.runs:
                        run.font.size = Pt(14)
                        heading.paragraph_format.space_after = Pt(2)
                        heading.paragraph_format.space_before = Pt(2)
            elif line.strip().startswith('- ') or line.strip().startswith('• '):
                # Bullet point
                bullet_text = line.strip()[2:]
                if use_emojis:
                    bullet_text = f"🔹 {bullet_text}"
                bullet_point = document.add_paragraph(style='List Bullet')
                run = bullet_point.add_run(bullet_text)
                run.font.size = base_font_size
                if compact_layout:
                    bullet_point.paragraph_format.space_after = Pt(2)
                    bullet_point.paragraph_format.line_spacing = 1.0
            elif line.strip():
                # Regular paragraph
                paragraph = document.add_paragraph(line.strip())
                if compact_layout:
                    for run in paragraph.runs:
                        run.font.size = base_font_size
                    paragraph.paragraph_format.space_after = Pt((1 if compact_layout else 6))
                    paragraph.paragraph_format.line_spacing = 1.0
            
            i += 1
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        logger.info(f"Word document generated: {output_path}")
        
        return output_path


def create_word_document(content: str, output_path: Path, title: str = "Processed Manual", compact_layout: bool = False, use_emojis: bool = False, remove_markdown_bold: bool = True, qr_code_path: Optional[Path] = None, qr_audio_url: Optional[str] = None) -> Path:
    """
    Convenience function to generate Word document
    """
    generator = WordGenerator()
    return generator.generate_word(content, output_path, title, compact_layout=compact_layout, use_emojis=use_emojis, remove_markdown_bold=remove_markdown_bold, qr_code_path=qr_code_path, qr_audio_url=qr_audio_url)